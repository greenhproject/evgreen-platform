from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

SOURCE = Path("/home/ubuntu/upload/Contrato_Aliado_Comercial_EVGreen_V2.docx")
OUTPUT = Path("/home/ubuntu/green-ev-platform/docs/Contrato_Aliado_EVGreen_V2_PLANTILLA_DINAMICA.docx")


def replace_paragraph_text(paragraph, new_text: str) -> None:
    """Sustituye el texto conservando el estilo principal del párrafo."""
    if paragraph.runs:
        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(new_text)


def replace_everywhere(document: Document, old: str, new: str) -> None:
    for paragraph in document.paragraphs:
        if old in paragraph.text:
            replace_paragraph_text(paragraph, paragraph.text.replace(old, new))
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if old in paragraph.text:
                        replace_paragraph_text(paragraph, paragraph.text.replace(old, new))


def set_table_value(table, row_index: int, marker: str) -> None:
    cell = table.rows[row_index].cells[1]
    paragraph = cell.paragraphs[0]
    replace_paragraph_text(paragraph, marker)


document = Document(SOURCE)

# Identidad del operador: se etiqueta para que una nueva razón social, contacto
# o representante pueda configurarse sin editar las cláusulas del contrato.
global_replacements = {
    "GREEN HOUSE PROJECT S.A.S.": "{{GHP_RAZON_SOCIAL}}",
    "GREEN HOUSE PROJECT S.A.S": "{{GHP_RAZON_SOCIAL}}",
    "901.447.678-0": "{{GHP_NIT}}",
    "Luis Enrique Narváez Rodríguez": "{{GHP_REPRESENTANTE}}",
    "1.015.418.125": "{{GHP_DOCUMENTO_REPRESENTANTE}}",
}
for original, marker in global_replacements.items():
    replace_everywhere(document, original, marker)

# Encabezado y versión controlada.
for paragraph in document.paragraphs:
    text = " ".join(paragraph.text.split())
    if text.startswith("VERSIÓN 2.0"):
        replace_paragraph_text(paragraph, "VERSIÓN {{VERSION_PLANTILLA}}")
    elif text.startswith("En señal de aceptación, las partes firman"):
        replace_paragraph_text(
            paragraph,
            "En señal de aceptación, las partes firman el presente contrato en {{CIUDAD_FIRMA}}, el {{FECHA_FIRMA}}.",
        )
    elif text.startswith("Nombre: _________________________________________"):
        replace_paragraph_text(paragraph, "Nombre: {{ALIADO_REPRESENTANTE}}")
    elif text.startswith("Cédula / NIT: _______________________________________"):
        replace_paragraph_text(paragraph, "Cédula / NIT: {{ALIADO_DOCUMENTO_REPRESENTANTE}}")
    elif text.startswith("Cargo (si aplica): __________________________"):
        replace_paragraph_text(paragraph, "Cargo: {{ALIADO_CARGO_REPRESENTANTE}}")
    elif text.startswith("Anexo 1: Plano y descripción del ESPACIO CEDIDO"):
        replace_paragraph_text(
            paragraph,
            "Anexo 1: Plano y descripción del ESPACIO CEDIDO — {{SITIO_NOMBRE}}, ubicado en {{SITIO_DIRECCION}}, "
            "{{SITIO_CIUDAD}}, {{SITIO_DEPARTAMENTO}}; área cedida: {{AREA_CEDIDA_M2}} m²; "
            "puestos destinados: {{PUESTOS_PARQUEO}}; plano/ficha anexa: {{PLANO_ANEXO_URL}}.",
        )
    elif text.startswith("9.1. Porcentaje de participación."):
        replace_paragraph_text(
            paragraph,
            "9.1. Porcentaje de participación. Como contraprestación por la cesión del espacio, EL OPERADOR pagará "
            "a EL ALIADO COMERCIAL una participación equivalente al {{PARTICIPACION_ALIADO_PORCENTAJE}}% del margen "
            "bruto mensual generado por la operación de la electrolinera instalada en el ESPACIO CEDIDO.",
        )
    elif text.startswith("b) Fecha de corte:"):
        replace_paragraph_text(paragraph, "b) Fecha de corte: {{FECHA_CIERRE_LIQUIDACION}}.")
    elif text.startswith("c) Fecha de pago:"):
        replace_paragraph_text(
            paragraph,
            "c) Fecha de pago: Dentro de los primeros {{PLAZO_PAGO_DIAS_HABILES}} días hábiles del mes siguiente "
            "al período liquidado.",
        )
    elif text.startswith("12.1. Vigencia."):
        replace_paragraph_text(
            paragraph,
            "12.1. Vigencia. El presente contrato tendrá una vigencia de {{PLAZO_INICIAL_ANOS}} años contados a partir "
            "de la fecha de puesta en operación comercial de la electrolinera, la cual será certificada por EL OPERADOR "
            "mediante acta de inicio de operación. La vigencia de {{PLAZO_INICIAL_ANOS}} años responde a la necesidad de "
            "garantizar la recuperación de la inversión realizada íntegramente por EL OPERADOR y la generación de retornos "
            "razonables para los inversionistas que financian la infraestructura de carga.",
        )
    elif text.startswith("12.2. Prórroga."):
        replace_paragraph_text(
            paragraph,
            "12.2. Prórroga. El contrato se prorrogará automáticamente por períodos sucesivos de {{PRORROGA_ANOS}} años, "
            "salvo que alguna de las partes notifique a la otra su intención de no prorrogar con al menos doce (12) meses "
            "de anticipación a la fecha de vencimiento del período vigente. Dentro de los seis (6) meses anteriores al "
            "vencimiento de cada período, las partes podrán solicitar una revisión integral de las condiciones del contrato "
            "que incluya: a) desempeño económico de la operación; b) condiciones del mercado de carga eléctrica; "
            "c) comportamiento de la demanda; y d) condiciones regulatorias aplicables. Dicha revisión no condicionará la "
            "prórroga automática, pero permitirá a las partes acordar ajustes mediante otrosí.",
        )

# Datos corporativos del operador (tabla 0) y de la EDS/aliado (tabla 1).
operator_values = {
    1: "{{GHP_RAZON_SOCIAL}}",
    2: "{{GHP_NIT}}",
    3: "{{GHP_REPRESENTANTE}}",
    4: "{{GHP_DOCUMENTO_REPRESENTANTE}}",
    5: "{{GHP_DOMICILIO}}",
    6: "{{GHP_CORREO_NOTIFICACIONES}}",
    7: "{{GHP_TELEFONO}}",
    8: "{{GHP_DIRECCION}}",
    9: "{{MARCA_COMERCIAL}}",
}
ally_values = {
    1: "{{ALIADO_RAZON_SOCIAL}}",
    2: "{{ALIADO_NIT}}",
    3: "{{ALIADO_REPRESENTANTE}}",
    4: "{{ALIADO_DOCUMENTO_REPRESENTANTE}}",
    5: "{{ALIADO_DOMICILIO}}",
    6: "{{ALIADO_CORREO_NOTIFICACIONES}}",
    7: "{{ALIADO_TELEFONO}}",
    8: "{{ALIADO_DIRECCION_NOTIFICACIONES}}",
}
for row_index, marker in operator_values.items():
    set_table_value(document.tables[0], row_index, marker)
for row_index, marker in ally_values.items():
    set_table_value(document.tables[1], row_index, marker)

# La sección de firmas del documento fuente está compuesta por párrafos al
# final. Se conservan la declaración de aceptación y las cláusulas previas,
# pero se retiran las líneas heredadas para que PDF manual y DocuSign inserten
# un único bloque de firma en el mismo punto lógico.
acceptance_found = False
for paragraph in list(document.paragraphs):
    normalized = " ".join(paragraph.text.split())
    if normalized.startswith("En señal de aceptación, las partes firman"):
        acceptance_found = True
        continue
    if acceptance_found:
        paragraph._element.getparent().remove(paragraph._element)

if not acceptance_found:
    raise RuntimeError("No se encontró la sección final de aceptación y firmas.")

signature_anchor = document.add_paragraph()
signature_anchor.alignment = WD_ALIGN_PARAGRAPH.CENTER
signature_anchor.add_run("EVG_SIGNATURE_BLOCK_HERE")

document.save(OUTPUT)
print(f"Plantilla dinámica creada: {OUTPUT}")
