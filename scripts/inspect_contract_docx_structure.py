from pathlib import Path
from docx import Document

SOURCE = Path("/home/ubuntu/upload/Contrato_Aliado_Comercial_EVGreen_V2.docx")
OUTPUT = Path("/home/ubuntu/green-ev-platform/docs/contrato-v2-estructura-docx.txt")

document = Document(SOURCE)
lines: list[str] = []

for index, paragraph in enumerate(document.paragraphs):
    text = " ".join(paragraph.text.split())
    if text:
        lines.append(f"P{index:04d}: {text}")

for table_index, table in enumerate(document.tables):
    lines.append(f"\nTABLE {table_index}")
    for row_index, row in enumerate(table.rows):
        cells = [" ".join(cell.text.split()) for cell in row.cells]
        lines.append(f"R{row_index:03d}: " + " || ".join(cells))

OUTPUT.write_text("\n".join(lines), encoding="utf-8")
print(f"Estructura extraída en {OUTPUT}")
