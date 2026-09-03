"""Extrae párrafos y tablas de una plantilla DOCX en un inventario legible."""

from pathlib import Path

from docx import Document


SOURCE = Path("/home/ubuntu/Downloads/contrato_aliado_evgreen_v2.docx")
OUTPUT = Path("/home/ubuntu/green-ev-platform/docs/contrato-aliado-evgreen-v2-extracto.md")


def main() -> None:
    document = Document(SOURCE)
    output: list[str] = ["# Extracto de plantilla contractual", ""]
    for index, paragraph in enumerate(document.paragraphs, start=1):
        text = paragraph.text.strip()
        if text:
            output.append(f"{index}. {text}")
    for table_index, table in enumerate(document.tables, start=1):
        output.extend(["", f"## Tabla {table_index}", ""])
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            output.append(" | ".join(cells))
    OUTPUT.write_text("\n".join(output) + "\n", encoding="utf-8")


if __name__ == "__main__":
    main()
